# cv_generator.py
import os
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List
from xml.sax.saxutils import escape as xml_escape

from docxtpl import DocxTemplate
from docx import Document
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import HRFlowable, Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"
OUTPUT_DIR = BASE_DIR / "output"
LOGO_CANDIDATES = [
    BASE_DIR / "assets" / "ntt_data_logo.png",
    BASE_DIR / "assets" / "NTT_Logo.png",
    BASE_DIR / "assets" / "ntt_logo.png",
    BASE_DIR / "assets" / "logo.png",
]

CV_LABELS = {
    "english": {
        "career_objective": "Career Objective",
        "professional_summary": "Professional Summary",
        "core_skills": "Core Skills",
        "education": "Education",
        "certifications": "Certifications",
        "achievements": "Achievements",
        "professional_experience": "Professional Experience",
        "profile_snapshot": "Profile Snapshot",
        "certifications_wins": "Certifications & Wins",
        "experience_highlights": "Experience Highlights",
        "summary": "Summary",
        "skills": "Skills",
        "experience": "Experience",
        "at": "at",
        "role": "Role",
        "company": "Company",
        "present": "Present",
        "summary_pending": "Summary pending.",
        "experience_pending": "Experience will appear here.",
        "education_pending": "Education pending.",
        "total_it_experience": "Total IT Experience",
    },
    "german": {
        "career_objective": "Karriereziel",
        "professional_summary": "Berufliches Profil",
        "core_skills": "Kernkompetenzen",
        "education": "Ausbildung",
        "certifications": "Zertifizierungen",
        "achievements": "Erfolge",
        "professional_experience": "Berufserfahrung",
        "profile_snapshot": "Profilubersicht",
        "certifications_wins": "Zertifizierungen und Erfolge",
        "experience_highlights": "Erfahrungsschwerpunkte",
        "summary": "Zusammenfassung",
        "skills": "Kenntnisse",
        "experience": "Erfahrung",
        "at": "bei",
        "role": "Rolle",
        "company": "Unternehmen",
        "present": "Heute",
        "summary_pending": "Zusammenfassung ausstehend.",
        "experience_pending": "Erfahrung erscheint hier.",
        "education_pending": "Ausbildung ausstehend.",
        "total_it_experience": "Gesamte IT-Erfahrung",
    },
    "spanish": {
        "career_objective": "Objetivo profesional",
        "professional_summary": "Resumen profesional",
        "core_skills": "Habilidades clave",
        "education": "Educacion",
        "certifications": "Certificaciones",
        "achievements": "Logros",
        "professional_experience": "Experiencia profesional",
        "profile_snapshot": "Resumen del perfil",
        "certifications_wins": "Certificaciones y logros",
        "experience_highlights": "Aspectos destacados de la experiencia",
        "summary": "Resumen",
        "skills": "Habilidades",
        "experience": "Experiencia",
        "at": "en",
        "role": "Rol",
        "company": "Empresa",
        "present": "Actualidad",
        "summary_pending": "Resumen pendiente.",
        "experience_pending": "La experiencia aparecera aqui.",
        "education_pending": "Educacion pendiente.",
        "total_it_experience": "Experiencia total en TI",
    },
    "hindi": {
        "career_objective": "Career Objective",
        "professional_summary": "Professional Summary",
        "core_skills": "Core Skills",
        "education": "Education",
        "certifications": "Certifications",
        "achievements": "Achievements",
        "professional_experience": "Professional Experience",
        "profile_snapshot": "Profile Snapshot",
        "certifications_wins": "Certifications & Wins",
        "experience_highlights": "Experience Highlights",
        "summary": "Summary",
        "skills": "Skills",
        "experience": "Experience",
        "at": "at",
        "role": "Role",
        "company": "Company",
        "present": "Present",
        "summary_pending": "Summary pending.",
        "experience_pending": "Experience will appear here.",
        "education_pending": "Education pending.",
        "total_it_experience": "Total IT Experience",
    },
    "japanese": {
        "career_objective": "キャリア目標",
        "professional_summary": "職務要約",
        "core_skills": "主要スキル",
        "education": "学歴",
        "certifications": "資格",
        "achievements": "実績",
        "professional_experience": "職務経歴",
        "profile_snapshot": "プロフィール概要",
        "certifications_wins": "資格と実績",
        "experience_highlights": "経験ハイライト",
        "summary": "要約",
        "skills": "スキル",
        "experience": "経験",
        "at": "at",
        "role": "役割",
        "company": "会社",
        "present": "現在",
        "summary_pending": "要約は未入力です。",
        "experience_pending": "経験はここに表示されます。",
        "education_pending": "学歴は未入力です。",
        "total_it_experience": "総IT経験",
    },
}


def _get_footer_logo():
    for logo_path in LOGO_CANDIDATES:
        if logo_path.exists():
            return ImageReader(str(logo_path))
    return None


def _draw_confidential_footer(canvas, doc):
    canvas.saveState()
    logo = _get_footer_logo()
    footer_y = 0.35 * inch
    if logo is not None:
        canvas.drawImage(
            logo,
            doc.leftMargin,
            footer_y - 0.08 * inch,
            width=0.9 * inch,
            height=0.26 * inch,
            preserveAspectRatio=True,
            mask="auto",
        )
    canvas.setFont("Helvetica-Bold", 9)
    canvas.setFillColor(colors.HexColor("#64748B"))
    canvas.drawCentredString(A4[0] / 2, footer_y, "NTT DATA CONFIDENTIAL")
    canvas.restoreState()


def _safe_filename(value: str, fallback: str = "candidate_cv") -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", (value or "").strip()).strip("_")
    return cleaned or fallback


def _normalize_output_language(output_language: str | None) -> str:
    value = (output_language or "English").strip().lower()
    aliases = {
        "en": "english",
        "english": "english",
        "de": "german",
        "german": "german",
        "deutsch": "german",
        "es": "spanish",
        "spanish": "spanish",
        "espanol": "spanish",
        "español": "spanish",
        "hi": "hindi",
        "hindi": "hindi",
        "ja": "japanese",
        "japanese": "japanese",
    }
    return aliases.get(value, "english")


def _get_cv_labels(output_language: str | None) -> Dict[str, str]:
    return CV_LABELS[_normalize_output_language(output_language)]


def _get_pdf_font_names(output_language: str | None) -> Dict[str, str]:
    if _normalize_output_language(output_language) == "japanese":
        try:
            pdfmetrics.getFont("HeiseiKakuGo-W5")
        except KeyError:
            pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
        return {
            "regular": "HeiseiKakuGo-W5",
            "bold": "HeiseiKakuGo-W5",
            "italic": "HeiseiKakuGo-W5",
        }
    return {
        "regular": "Helvetica",
        "bold": "Helvetica-Bold",
        "italic": "Helvetica-Oblique",
    }


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.text = new_text
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _localize_generated_docx(out_path: Path, labels: Dict[str, str]) -> None:
    document = Document(str(out_path))
    replacements = {
        "CAREER OBJECTIVE": labels["career_objective"].upper(),
        "PROFESSIONAL SUMMARY": labels["professional_summary"].upper(),
        "CORE SKILLS": labels["core_skills"].upper(),
        "EDUCATION": labels["education"].upper(),
        "CERTIFICATIONS": labels["certifications"].upper(),
        "ACHIEVEMENTS": labels["achievements"].upper(),
        "PROFESSIONAL EXPERIENCE": labels["professional_experience"].upper(),
    }

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            _replace_paragraph_text(paragraph, replacements[text])

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text in replacements:
                        _replace_paragraph_text(paragraph, replacements[text])

    document.save(str(out_path))


def _flatten_education(education_items: Iterable[Any]) -> List[str]:
    rows = []
    for item in education_items or []:
        if isinstance(item, dict):
            values = [str(value).strip() for value in item.values() if str(value).strip()]
            if values:
                rows.append(" | ".join(values))
        elif str(item).strip():
            rows.append(str(item).strip())
    return rows


def _normalize_experience(experience_items: Iterable[Any]) -> List[Dict[str, Any]]:
    rows = []
    for item in experience_items or []:
        if not isinstance(item, dict):
            continue
        normalized = dict(item)
        responsibilities = item.get("responsibilities") or item.get("bullets") or []
        if isinstance(responsibilities, str):
            responsibilities = [responsibilities] if responsibilities.strip() else []
        normalized["responsibilities"] = [str(value).strip() for value in responsibilities if str(value).strip()]
        rows.append(normalized)
    return rows


def _safe_text(value: Any) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, dict):
        parts = [str(item).strip() for item in value.values() if str(item).strip()]
        return " | ".join(parts)
    if isinstance(value, list):
        parts = [str(item).strip() for item in value if str(item).strip()]
        return ", ".join(parts)
    return str(value).strip()


def _clean_pdf_text(value: Any) -> str:
    text = _safe_text(value)
    if not text:
        return ""
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", " ", text)


def _xml_text(value: Any) -> str:
    return xml_escape(_clean_pdf_text(value))


def _build_profile_photo(profile_photo_bytes: Any, width: float, height: float):
    if not profile_photo_bytes:
        return None
    try:
        photo = RLImage(BytesIO(profile_photo_bytes), width=width, height=height)
        photo.hAlign = "RIGHT"
        return photo
    except Exception:
        return None


def generate_docx(
    cv_data: Dict[str, Any],
    output_filename: str = "generated_cv.docx",
    template_id: str = "custom",
    output_language: str = "English",
) -> str:
    """
    Fill the Word template placeholders and save a DOCX copy.
    """
    template_path = TEMPLATES_DIR / "cv_template_clean.docx"
    if not template_path.exists():
        template_path = TEMPLATES_DIR / "cv_template.docx"
    doc = DocxTemplate(str(template_path))

    context = {
        "objectives": cv_data.get("objectives", ""),
        "name": cv_data.get("name", ""),
        "title": cv_data.get("title", ""),
        "total_it_experience": cv_data.get("total_it_experience", ""),
        "contact": cv_data.get("contact", ""),
        "location": cv_data.get("location", ""),
        "summary": cv_data.get("summary", ""),
        "skills": cv_data.get("skills", []),
        "education": _flatten_education(cv_data.get("education", [])),
        "experience": _normalize_experience(cv_data.get("experience", [])),
        "certifications": cv_data.get("certifications", []),
        "achievements": cv_data.get("achievements", []),
        "template_id": template_id,
    }

    doc.render(context)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out_path = OUTPUT_DIR / output_filename
    doc.save(str(out_path))
    _localize_generated_docx(out_path, _get_cv_labels(output_language))
    return str(out_path)


def _build_common_pdf_styles(output_language: str | None = None):
    fonts = _get_pdf_font_names(output_language)
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ResumeHeader",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontName=fonts["bold"],
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#0F172A"),
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeSubheader",
            parent=styles["BodyText"],
            alignment=TA_CENTER,
            fontName=fonts["regular"],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#475569"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionHeading",
            parent=styles["Heading2"],
            fontName=fonts["bold"],
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#0F766E"),
            spaceBefore=12,
            spaceAfter=6,
            uppercase=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="RoleHeading",
            parent=styles["BodyText"],
            fontName=fonts["bold"],
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#0F172A"),
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="MetaText",
            parent=styles["BodyText"],
            fontName=fonts["italic"],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#64748B"),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletText",
            parent=styles["BodyText"],
            fontName=fonts["regular"],
            fontSize=10,
            leading=14,
            leftIndent=12,
            bulletIndent=0,
            textColor=colors.HexColor("#1E293B"),
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CardHeader",
            parent=styles["Heading2"],
            fontName=fonts["bold"],
            fontSize=12,
            leading=15,
            textColor=colors.white,
            alignment=TA_LEFT,
        )
    )
    return styles


def _create_output_path(cv_data: Dict[str, Any], output_filename: str | None) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    candidate_name = cv_data.get("name", "")
    filename = output_filename or f"{_safe_filename(candidate_name)}.pdf"
    return str(OUTPUT_DIR / filename)


def _build_custom_story(cv_data: Dict[str, Any], styles, labels: Dict[str, str]):
    story = []
    name = cv_data.get("name") or "Candidate Name"
    title = cv_data.get("title") or "Professional Title"
    experience_text = cv_data.get("total_it_experience") or ""
    location = cv_data.get("location") or ""
    contact = cv_data.get("contact") or ""
    summary = (cv_data.get("summary") or "").strip()
    objectives = (cv_data.get("objectives") or "").strip()
    profile_photo_bytes = cv_data.get("profile_photo_bytes")

    subheader_parts = [part for part in [title, experience_text, location, contact] if part]

    photo = _build_profile_photo(profile_photo_bytes, 1.1 * inch, 1.35 * inch)
    if photo:
        header_flowables = [Paragraph(name, styles["ResumeHeader"])]
        if subheader_parts:
            header_flowables.append(Paragraph(" | ".join(subheader_parts), styles["ResumeSubheader"]))
        header_table = Table(
            [[header_flowables, photo]],
            colWidths=[5.65 * inch, 1.0 * inch],
            hAlign="LEFT",
        )
        header_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ]
            )
        )
        story.append(header_table)
    else:
        story.append(Paragraph(name, styles["ResumeHeader"]))
        if subheader_parts:
            story.append(Paragraph(" | ".join(subheader_parts), styles["ResumeSubheader"]))

    story.append(HRFlowable(width="100%", thickness=1.2, color=colors.HexColor("#CBD5E1")))
    story.append(Spacer(1, 0.12 * inch))

    def add_section(title_text: str, body_items: List[Any]):
        if not body_items:
            return
        story.append(Paragraph(title_text, styles["SectionHeading"]))
        story.extend(body_items)

    if objectives:
        add_section(labels["career_objective"], [Paragraph(objectives, styles["BodyText"])])

    summary_blocks = []
    if experience_text:
        summary_blocks.append(Paragraph(f"{labels['total_it_experience']}: {experience_text}", styles["BodyText"]))
    if summary:
        summary_blocks.append(Paragraph(summary, styles["BodyText"]))
    add_section(labels["professional_summary"], summary_blocks)

    skills = [str(skill).strip() for skill in cv_data.get("skills", []) if str(skill).strip()]
    if skills:
        add_section(labels["core_skills"], [Paragraph(" | ".join(skills), styles["BodyText"])])

    experience_blocks = []
    for item in _normalize_experience(cv_data.get("experience", [])):
        role = item.get("role") or labels["role"]
        company = item.get("company") or labels["company"]
        start_date = item.get("start_date") or ""
        end_date = item.get("end_date") or labels["present"]

        experience_blocks.append(Paragraph(f"{role} | {company}", styles["RoleHeading"]))
        if start_date or end_date:
            experience_blocks.append(Paragraph(f"{start_date} - {end_date}".strip(" -"), styles["MetaText"]))
        for responsibility in item.get("responsibilities", []):
            experience_blocks.append(Paragraph(responsibility, styles["BulletText"], bulletText="-"))
        experience_blocks.append(Spacer(1, 0.04 * inch))
    add_section(labels["professional_experience"], experience_blocks)

    education_rows = [Paragraph(item, styles["BodyText"]) for item in _flatten_education(cv_data.get("education", []))]
    add_section(labels["education"], education_rows)

    certification_rows = [
        Paragraph(str(item).strip(), styles["BulletText"], bulletText="-")
        for item in cv_data.get("certifications", [])
        if str(item).strip()
    ]
    add_section(labels["certifications"], certification_rows)

    achievement_rows = [
        Paragraph(str(item).strip(), styles["BulletText"], bulletText="-")
        for item in cv_data.get("achievements", [])
        if str(item).strip()
    ]
    add_section(labels["achievements"], achievement_rows)
    return story


def _card(title: str, body: List[Any], width: float, accent_hex: str, output_language: str | None = None):
    fonts = _get_pdf_font_names(output_language)
    header = Table(
        [[Paragraph(_xml_text(title), ParagraphStyle("CardHeaderInline", fontName=fonts["bold"], fontSize=11, textColor=colors.white))]],
        colWidths=[width],
    )
    header.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(accent_hex)),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    rows = [[item] for item in body] if body else [[Spacer(1, 0.01 * inch)]]
    body_table = Table(rows, colWidths=[width], splitByRow=1)
    body_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
                ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#D6E3F0")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return [header, body_table]


def _build_postcard_story(cv_data: Dict[str, Any], styles, labels: Dict[str, str], output_language: str | None = None):
    fonts = _get_pdf_font_names(output_language)
    name = _safe_text(cv_data.get("name")) or "Candidate Name"
    title = _safe_text(cv_data.get("title")) or "Professional Title"
    summary = _safe_text(cv_data.get("summary")) or labels["summary_pending"]
    profile_photo_bytes = cv_data.get("profile_photo_bytes")
    skills = [_safe_text(skill) for skill in cv_data.get("skills", []) if _safe_text(skill)]
    certifications = [_safe_text(item) for item in cv_data.get("certifications", []) if _safe_text(item)]
    achievements = [_safe_text(item) for item in cv_data.get("achievements", []) if _safe_text(item)]
    education = _flatten_education(cv_data.get("education", []))
    experience = _normalize_experience(cv_data.get("experience", []))

    hero_copy = Paragraph(
        (
            f"<font size='22'><b>{_xml_text(name)}</b></font>"
            f"<br/><font size='11'>{_xml_text(title)}</font>"
            f"<br/><font size='10'>{_xml_text(cv_data.get('total_it_experience'))}</font>"
        ),
        ParagraphStyle("PostcardHero", fontName=fonts["regular"], fontSize=11, leading=14, textColor=colors.white),
    )
    photo = _build_profile_photo(profile_photo_bytes, 0.95 * inch, 1.12 * inch)
    if photo:
        hero = Table([[hero_copy, photo]], colWidths=[5.85 * inch, 1.15 * inch])
    else:
        hero = Table([[hero_copy]], colWidths=[7.0 * inch])
    hero.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#0E1A3B")),
                ("LEFTPADDING", (0, 0), (-1, -1), 18),
                ("RIGHTPADDING", (0, 0), (-1, -1), 18),
                ("TOPPADDING", (0, 0), (-1, -1), 18),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 18),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )

    left_body = [Paragraph(_xml_text(summary), styles["BodyText"])]
    if skills:
        left_body.append(Spacer(1, 0.08 * inch))
        left_body.append(Paragraph(_xml_text(" | ".join(skills)), styles["BodyText"]))
    right_body = []
    for row in certifications or ["No certifications listed"]:
        right_body.append(Paragraph(_xml_text(row), styles["BulletText"], bulletText="-"))
    if achievements:
        right_body.append(Spacer(1, 0.08 * inch))
        for row in achievements:
            right_body.append(Paragraph(_xml_text(row), styles["BulletText"], bulletText="-"))

    card_width = 6.95 * inch
    cards = [hero, Spacer(1, 0.16 * inch)]
    cards.extend(_card(labels["profile_snapshot"], left_body, card_width, "#1479C9", output_language))
    cards.append(Spacer(1, 0.14 * inch))
    cards.extend(_card(labels["certifications_wins"], right_body, card_width, "#0E1A3B", output_language))
    cards.append(Spacer(1, 0.14 * inch))

    experience_rows: List[Any] = []
    for item in experience:
        role = item.get("role") or labels["role"]
        company = item.get("company") or labels["company"]
        experience_rows.append(Paragraph(f"<b>{_xml_text(role)}</b> | {_xml_text(company)}", styles["BodyText"]))
        for bullet in item.get("responsibilities", [])[:3]:
            experience_rows.append(Paragraph(_xml_text(bullet), styles["BulletText"], bulletText="-"))
        experience_rows.append(Spacer(1, 0.05 * inch))
    if not experience_rows:
        experience_rows.append(Paragraph(_xml_text(labels["experience_pending"]), styles["BodyText"]))

    education_rows = [Paragraph(_xml_text(item), styles["BodyText"]) for item in education] or [Paragraph(_xml_text(labels["education_pending"]), styles["BodyText"])]
    cards.extend(_card(labels["experience_highlights"], experience_rows, card_width, "#18A0AE", output_language))
    cards.append(Spacer(1, 0.14 * inch))
    cards.extend(_card(labels["education"], education_rows, card_width, "#1479C9", output_language))
    return cards


def _build_sample_story(cv_data: Dict[str, Any], styles, labels: Dict[str, str]):
    story: List[Any] = []
    name = cv_data.get("name") or "Candidate Name"
    title = cv_data.get("title") or "Professional Title"
    meta_parts = [part for part in [cv_data.get("location"), cv_data.get("contact"), cv_data.get("total_it_experience")] if part]
    story.append(Paragraph(name, ParagraphStyle("SampleName", fontName="Helvetica-Bold", fontSize=24, textColor=colors.HexColor("#102542"), spaceAfter=4)))
    story.append(Paragraph(title, ParagraphStyle("SampleTitle", fontName="Helvetica", fontSize=12, textColor=colors.HexColor("#1479C9"), spaceAfter=5)))
    if meta_parts:
        story.append(Paragraph(" | ".join(meta_parts), styles["MetaText"]))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1479C9")))
    story.append(Spacer(1, 0.12 * inch))

    summary = (cv_data.get("summary") or "").strip()
    if summary:
        story.append(Paragraph(labels["summary"], styles["SectionHeading"]))
        story.append(Paragraph(summary, styles["BodyText"]))

    skills = [str(skill).strip() for skill in cv_data.get("skills", []) if str(skill).strip()]
    if skills:
        story.append(Paragraph(labels["skills"], styles["SectionHeading"]))
        story.append(Paragraph("  •  ".join(skills), styles["BodyText"]))

    story.append(Paragraph(labels["experience"], styles["SectionHeading"]))
    experience = _normalize_experience(cv_data.get("experience", []))
    if experience:
        for item in experience:
            role = item.get("role") or labels["role"]
            company = item.get("company") or labels["company"]
            dates = " - ".join(filter(None, [item.get("start_date"), item.get("end_date")]))
            story.append(Paragraph(f"<b>{role}</b> {labels['at']} {company}", styles["BodyText"]))
            if dates:
                story.append(Paragraph(dates, styles["MetaText"]))
            for bullet in item.get("responsibilities", []):
                story.append(Paragraph(bullet, styles["BulletText"], bulletText="-"))
            story.append(Spacer(1, 0.05 * inch))
    else:
        story.append(Paragraph(labels["experience_pending"], styles["BodyText"]))

    education = _flatten_education(cv_data.get("education", []))
    if education:
        story.append(Paragraph(labels["education"], styles["SectionHeading"]))
        for item in education:
            story.append(Paragraph(item, styles["BodyText"]))

    certifications = [str(item).strip() for item in cv_data.get("certifications", []) if str(item).strip()]
    if certifications:
        story.append(Paragraph(labels["certifications"], styles["SectionHeading"]))
        for item in certifications:
            story.append(Paragraph(item, styles["BulletText"], bulletText="-"))

    return story


def generate_pdf(cv_data: Dict[str, Any], output_filename: str | None = None, template_id: str = "custom", output_language: str = "English") -> str:
    out_path = _create_output_path(cv_data, output_filename)
    styles = _build_common_pdf_styles(output_language)
    labels = _get_cv_labels(output_language)
    name = cv_data.get("name") or "Candidate Name"

    if template_id == "postcard":
        story = _build_postcard_story(cv_data, styles, labels, output_language)
    elif template_id == "sample":
        story = _build_sample_story(cv_data, styles, labels)
    else:
        story = _build_custom_story(cv_data, styles, labels)

    doc = SimpleDocTemplate(
        out_path,
        pagesize=A4,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.55 * inch,
        title=f"{name} CV",
        author=name,
    )
    doc.build(story, onFirstPage=_draw_confidential_footer, onLaterPages=_draw_confidential_footer)
    return out_path
